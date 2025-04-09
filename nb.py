import asyncio
import threading

import logging
from aiogram import Bot, Dispatcher, types
from aiogram.filters.command import Command
from aiogram.exceptions import TelegramBadRequest
from aiogram.types import ReplyKeyboardRemove, \
    ReplyKeyboardMarkup, KeyboardButton, \
    InlineKeyboardMarkup, InlineKeyboardButton
import os
import random
from aiogram.types.input_file import FSInputFile
from aiogram.types import InputFile
from templates.keyboards import *
from cache_cleaner import cache_clean
from audio_utils import transcoding_to_mp3
from json_utils import Json_work
from ai_utils import GPTClient
from mongo_utils import DBClient
from aiogram.client.telegram import TelegramAPIServer
import moviepy.editor as mp
from aiogram.client.session.aiohttp import AiohttpSession
from config import *
from aiogram.client.telegram import TelegramAPIServer
import asyncio
from concurrent.futures import ThreadPoolExecutor
from moviepy.editor import VideoFileClip

import docx

#from bot_templates.keyboards import keyboard, keyboard_audio
J = Json_work()
gpts = []
for ass in asses:
    gpts.append(GPTClient(gpt_token, gpt_model, gpt_voice_model, ass))
DB = DBClient(mongo_client_ref, db_name)
bot_templates_path = {"keyboard":os.path.join("templates", "keyboard_temp.json"), "messages":os.path.join("templates", "msg_temp.json")}

logging.basicConfig(level=logging.INFO) 
threading.Thread(target=cache_clean).start()
logging.info(f"thread {cache_clean.__name__} is starting")
session = AiohttpSession(api=TelegramAPIServer.from_base("http://localhost:8081", is_local=True))
bot = Bot(token=bot_token, session=session)
dp = Dispatcher()
   

import aiohttp
import asyncio
import base64
from config import speech_secret_key, speech_token, bucket_name
import time
import json
out="https://stt.api.cloud.yandex.net/stt/v3/getRecognition" #+"<operation_id>"


class Bucket_tools:
    def __init__(self, api_key:str,  bucket_name:str):
        self.api_key = api_key
        self.bucket_name = bucket_name
    async def object_upload(self, object_path:str, local_file_path:str)->str: #object path - path in bucket to object
        request = f"https://storage.yandexcloud.net/{self.bucket_name}/{object_path}"
        headers = {
                    'Authorization': f'Api-key {self.api_key}',
                    'Content-Type': 'application/octet-stream',
        }
        encoded_data = open(local_file_path, "rb").read()
        async with aiohttp.ClientSession(headers=headers) as session:
            response = await session.put(request, data=encoded_data)
            return await response.text()

    async def object_get(self, object_path:str)->str:
        request = f"https://storage.yandexcloud.net/{self.bucket_name}/{object_path}"
        headers = {
                    'Authorization': f'Api-key {self.api_key}',
                    'Content-Type': 'application/octet-stream',
        }
        async with aiohttp.ClientSession(headers=headers) as session:
            response = await session.get(request)
            return await response.text()


class Speech_recognizer:
    def __init__(self, api_key: str, bucket_name: str):
        self.api_key = api_key
        self.bucket_name = bucket_name

    async def audio_up(self, bucket_aud_name: str) -> dict:
        url =  "https://stt.api.cloud.yandex.net/stt/v3/recognizeFileAsync"
        request = {
                    "uri": f"https://storage.yandexcloud.net/{self.bucket_name}/{bucket_aud_name}",
                    "recognitionModel": {
                    "model": "general",
                    "audioFormat": {
                        "containerAudio": {
                                    "containerAudioType": "MP3"
                    }
                }
            }
        }

        headers = {
            'Authorization': f'Api-key {self.api_key}',
            'Content-Type': 'application/json'
        }
        print(request)
        async with aiohttp.ClientSession() as session:
            async with session.post(url, json=request, headers=headers) as response:
                return await response.json()

    async def operation_get(self, operation_id: str) -> dict:
        url =  f"https://operation.api.cloud.yandex.net/operations/{operation_id}"
        headers = {
            'Authorization': f'Api-key {self.api_key}',
        }
        async with aiohttp.ClientSession() as session:
            async with session.get(url, headers=headers) as response:
                return await response.json()

    async def transcript_get(self, operation_id: str) -> dict:
        url =  f"https://stt.api.cloud.yandex.net/stt/v3/getRecognition?operationId={operation_id}"
        headers = {
            'Authorization': f'Api-key {self.api_key}',
        }
        async with aiohttp.ClientSession() as session:
            async with session.get(url, headers=headers) as response:
                return await response.text()

bt = Bucket_tools(speech_token, bucket_name)

sr = Speech_recognizer(speech_secret_key, bucket_name)

import collections
state = collections.defaultdict(int)
current = collections.defaultdict(str)
prev = collections.defaultdict(str)
user_thread = collections.defaultdict(str)
dop = collections.defaultdict(str)
band = collections.defaultdict(lambda: True)

from formatting import markdown_to_html, escape_symbols, escape_markdown_symbols

from config import reply_chat_id
from aiogram.enums.parse_mode import ParseMode

MAX_MESSAGE_LENGTH = 4096 - 10
SPECIAL_SYMBOLS = ["*", "_", "~"]

async def send_big_message(bot, user_id, text, reply_markup=None):
    if not text:
        return
    symbols_stack = []
    code_mode = False
    big_code_mode = False
    text = escape_markdown_symbols(text)
    for i in range(0, len(text), MAX_MESSAGE_LENGTH):
        text_part = ""
        j = 0
        if big_code_mode:
            text_part += "```"
            j += 3
        elif code_mode:
            text_part += "`"
            j += 1
        else:
            for symbol in symbols_stack:
                text_part += symbol
        text_part += text[i:i + MAX_MESSAGE_LENGTH]
        while j < len(text_part):
            if j + 3 <= len(text_part) and text_part[j:j + 3] == "```":
                big_code_mode = not big_code_mode
                j += 3
                continue
            if text_part[j] == "`":
                code_mode = not code_mode
                j += 1
                continue
            if not code_mode and j + 2 <= len(text_part) and text_part[j:j + 2] == "__":
                if symbols_stack and symbols_stack[-1] == "__":
                    symbols_stack.pop()
                else:
                    symbols_stack.append("__")
                j += 2
                continue
            if not code_mode and not big_code_mode and text_part[j] in SPECIAL_SYMBOLS:
                if symbols_stack and symbols_stack[-1] == text_part[j]:
                    symbols_stack.pop()
                else:
                    symbols_stack.append(text_part[j])
            j += 1
        if big_code_mode:
            text_part += "```"
        elif code_mode:
            text_part += "`"
        else:
            for symbol in symbols_stack[::-1]:
                text_part += symbol
        try:
            await bot.send_message(user_id, text_part, parse_mode=ParseMode.MARKDOWN_V2, reply_markup=reply_markup)
            continue
        except Exception as e:
            await bot.send_message(reply_chat_id, e)
            print(e)
        for parse_mode in [ParseMode.MARKDOWN, ParseMode.HTML]:
            try:
                await bot.send_message(user_id, text_part, parse_mode=parse_mode, reply_markup=reply_markup)
                break
            except Exception as e:
                await bot.send_message(reply_chat_id, e)
                print(e)
        else:
            try:
                await bot.send_message(user_id, markdown_to_html(text_part), parse_mode=ParseMode.HTML, reply_markup=reply_markup)
                continue
            except Exception as e:
                await bot.send_message(reply_chat_id, e)
                print(e)
            try:
                await bot.send_message(user_id, escape_symbols(text_part), reply_markup=reply_markup)
                continue
            except Exception as e:
                await bot.send_message(reply_chat_id, e)
                print(e)


@dp.message(Command("start"))
async def cmd_start(message: types.Message):
#     id_check(message.from_user.id)
    state[message.from_user.id] = 0
    dop[message.from_user.id] = ""
    current[message.from_user.id] = ""
    prev[message.from_user.id] = ""
    await message.answer(J.read_p(bot_templates_path["messages"])["start"], reply_markup=keyboard_start)

    await bot.forward_message(chat_id= reply_chat_id, from_chat_id=message.chat.id, message_id=message.message_id)
    await bot.forward_message(chat_id= reply_chat_id, from_chat_id=message.chat.id, message_id=message.message_id+1)


@dp.message(Command("info"))
async def cmd_info(message):
    with open("promts/dop_info.txt") as f:
        lines = f.readlines()
        await send_big_message(bot, message.from_user.id, "\n".join(lines), reply_markup=no_keyboard)
    await bot.forward_message(chat_id= reply_chat_id, from_chat_id=message.chat.id, message_id=message.message_id)
    await bot.forward_message(chat_id= reply_chat_id, from_chat_id=message.chat.id, message_id=message.message_id+1)
    state[message.from_user.id] = -1


@dp.message(Command("band"))
async def cmd_band(message):
    if band[message.from_user.id]:
        await message.answer(f"Сейчас оркестр ассистентов включен. Выберите нужный вариант взаимодействия с ботом, нажав одну из кнопок.", reply_markup=keyboard_band)
    else:
        await message.answer(f"Сейчас оркестр ассистентов отключен. Выберите нужный вариант взаимодействия с ботом, нажав одну из кнопок.", reply_markup=keyboard_band)
    await bot.forward_message(chat_id= reply_chat_id, from_chat_id=message.chat.id, message_id=message.message_id)
    await bot.forward_message(chat_id= reply_chat_id, from_chat_id=message.chat.id, message_id=message.message_id+1)


@dp.message(Command("doc"))
async def cmd_doc(message: types.Message):
    # Получаем предыдущее сообщение
    previous_message = prev[message.from_user.id]
    if isinstance(previous_message, list):
        previous_message = "\n\n".join(previous_message)

    # Создаем новый документ
    doc = docx.Document()
    doc.add_paragraph(previous_message)

    # Сохраняем документ в файл
    doc_filename = f"cache/chapter{random.randint(10, 100)}.docx"
    doc.save(doc_filename)

    # Отправляем файл пользователю
    await message.answer_document(FSInputFile(path=doc_filename))


@dp.message(Command("intro"))
async def cmd_intro(message):
    await message.answer("Отправьте документ, для которого нужно сформировать введение", reply_markup=no_keyboard)
    await bot.forward_message(chat_id= reply_chat_id, from_chat_id=message.chat.id, message_id=message.message_id)
    await bot.forward_message(chat_id= reply_chat_id, from_chat_id=message.chat.id, message_id=message.message_id+1)
    state[message.from_user.id] = -2


@dp.message(Command("concl"))
async def cmd_conclusion(message):
    await message.answer("Отправьте документ, для которого нужно сформировать заключение", reply_markup=no_keyboard)
    await bot.forward_message(chat_id= reply_chat_id, from_chat_id=message.chat.id, message_id=message.message_id)
    await bot.forward_message(chat_id= reply_chat_id, from_chat_id=message.chat.id, message_id=message.message_id+1)
    state[message.from_user.id] = -1


@dp.message(Command("metod"))
async def cmd_metod(message):
    await message.answer("Отправьте документ, для которого нужно сформировать методику", reply_markup=no_keyboard)
    await bot.forward_message(chat_id= reply_chat_id, from_chat_id=message.chat.id, message_id=message.message_id)
    await bot.forward_message(chat_id= reply_chat_id, from_chat_id=message.chat.id, message_id=message.message_id+1)
    state[message.from_user.id] = -3


@dp.message(Command("continue"))
async def cmd_continue(message: types.Message):
    global current
    global user_thread

    if state[message.from_user.id] >= len(gpts) - 3:
        await message.answer(f"Написание главы завершено! Нажмите Заново, чтобы написать новую главу")
        return
    elif current[message.from_user.id]:
        gpt = gpts[state[message.from_user.id]]
        await message.answer("Работает ассистент " + ass_name[state[message.from_user.id]] + "...", reply_markup=no_keyboard)
        tr = gpt.create_thread()
        user_thread[message.from_user.id] = tr
        pr = prev[message.from_user.id]
        dp = dop[message.from_user.id]
        if state[message.from_user.id] == 3:
            topics = pr.split("###")
            if len(topics) < 2:
                topics = pr.split("\n")
            prev[message.from_user.id] = []
            for topic in topics[1:-1]:
                if not topic:
                    continue
                mess = (f"\nКак можно подробнее опиши эту часть главы: {topic}" if topic else "") + (f"\nДополнительная информация: {dp}" if dp else "") + "\nНе пиши название главы, только ее текст"
                gpt.add_message(tr, mess)
                answer = gpt.get_answer(tr)
                prev[message.from_user.id].append(answer)
                if band[message.from_user.id]:
                    await send_big_message(bot, message.from_user.id, answer, reply_markup=keyboard_continue)
        elif state[message.from_user.id] > 3:
            new_prev = []
            for pr in prev[message.from_user.id]:
                mess = (f"\nИсходный текст: {pr}" if pr else "") + (f"\nДополнительная информация: {dp}" if dp else "") + "\nНапиши только исправленный текст главы, без комментариев"
                gpt.add_message(tr, mess)
                answer = gpt.get_answer(tr)
                new_prev.append(answer)
                if band[message.from_user.id] or state[message.from_user.id] >= len(gpts) - 4:
                    await send_big_message(bot, message.from_user.id, answer, reply_markup=keyboard_continue)
            prev[message.from_user.id] = new_prev
        else:
            mess = f"Исходный текст: {current[message.from_user.id]}" + (f"\nАнализ текста: {pr}" if pr else "") + (f"\nДополнительная информация: {dp}" if dp else "")
            gpt.add_message(tr, mess)
            answer = gpt.get_answer(tr)
            prev[message.from_user.id] = answer
            if band[message.from_user.id]:
                await send_big_message(bot, message.from_user.id, answer, reply_markup=keyboard_continue)
    state[message.from_user.id] += 1

    if band[message.from_user.id]:   
        await bot.forward_message(chat_id= reply_chat_id, from_chat_id=message.chat.id, message_id=message.message_id)
        await bot.forward_message(chat_id= reply_chat_id, from_chat_id=message.chat.id, message_id=message.message_id+1)
        await bot.forward_message(chat_id= reply_chat_id, from_chat_id=message.chat.id, message_id=message.message_id+2)
        await send_big_message(bot, message.from_user.id, "Нажмите кнопку *Продолжить* для продолжения...", reply_markup=keyboard_continue)
    else:
        print("CMD CONTINUE!")
        await cmd_continue(message)


@dp.message(Command("assist"))
async def cmd_assist(message: types.Message):
    global state
    try:
        # Разделяем команду и аргумент, например, /assist 5
        parts = message.text.split()
        if len(parts) > 1:
            # Пытаемся преобразовать аргумент в число
            state[message.from_user.id] = int(parts[1])
            await message.reply(f"State has been set to {state[message.from_user.id]}.")
            await cmd_continue(message)
        else:
            await message.reply("Please provide a number after /assist. Example: /assist 5")
    except ValueError:
        await message.reply("Invalid number format. Please use an integer.")


def read_docx_to_string(file_path):
    # Load the document
    doc = docx.Document(file_path)

    # Extract text from each paragraph
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)

    # Join all paragraphs into a single string
    return '\n'.join(full_text)


async def convert_video_to_audio(file_path, output_path):
    loop = asyncio.get_event_loop()
    # Use a ThreadPoolExecutor to run the blocking I/O operation in a separate thread
    with ThreadPoolExecutor() as executor:
        await loop.run_in_executor(
            executor, extract_audio, file_path, output_path
        )


def extract_audio(file_path, output_path):
    # This function will run in a separate thread
    clip = VideoFileClip(file_path)
    print(min(MAX_DURATION * 60 * 60, clip.duration))
    clip = clip.subclip(0, min(MAX_DURATION * 60, clip.duration))
    clip.audio.write_audiofile(output_path)


@dp.message()
async def GPT_data(message: types.Message):

    async def process_audio(filename, file=None):
        fp=os.path.join(cache_path, filename)
        if file:
            downloaded_file = await bot.download_file(file.file_path)
            open(fp, 'wb').write(downloaded_file.read())
        await bt.object_upload(filename, fp)
        res = eval(str(await sr.audio_up(filename)))
        print(res)
        id = res["id"]
        print(id)
        async def rn():
            res = False
            while res != True:
                await asyncio.sleep(1)
                res = eval(str(await sr.operation_get(id)))["done"]
        await rn()

        result = await sr.transcript_get(id)

        # Разбиваем строку на отдельные JSON-объекты
        json_objects = result.strip().split('\n')

        # Список для хранения всех объединенных текстов из каждого JSON-объекта
        all_concatenated_texts = []
        for json_object in json_objects[::2]:
            try:
                data = json.loads(json_object)
                # Проверяем, существует ли ключ "final" в данном JSON-объекте
                if "final" in data['result']:
                    concatenated_text = ' '.join(
                        word['text'] for word in data['result']['final']['alternatives'][0]['words']
                    )
                    all_concatenated_texts.append(concatenated_text)
            except json.JSONDecodeError:
                await message.answer("ошибки декодирования JSON")
            except KeyError:
                await message.answer("ошибки ключей")

        # Выводим каждый объединенный текст из отдельных JSON-объектов
        global current, state
        current[message.from_user.id] = ""
        for text in all_concatenated_texts:
            if text:
                # await message.answer(text)
                current[message.from_user.id] += " " + text
        state[message.from_user.id] = 0

    async def process_video(file, strp, caption=""):
        try:
            if strp[1] in ('mp4', 'avi', 'mkv', 'mov', 'wmv'):
                await message.answer("Конвертирую видео в аудио...")
                await convert_video_to_audio(file.file_path, os.path.join(cache_path, f"{strp[0]}.mp3"))
                await message.answer("Расшифровываю аудио...")
                await process_audio(f"{strp[0]}.mp3")
                if caption:
                    current[message.from_user.id] += f"\nДополнительная информация: {caption}"
                if not current[message.from_user.id]:
                    await message.answer("Не удалось распознать текст")
                    return
                await cmd_continue(message)
            else:
                await message.answer("Извините, но данный формат видео не поддерживается.")
        except TelegramBadRequest as e:
            await message.answer(f"Произошла ошибка: {e}")

    tr = user_thread[message.from_user.id] or gpts[state[message.from_user.id]].create_thread()
    if message.text and message.chat.id != reply_chat_id:
        match message.text:
            case "Продолжить":
                await cmd_continue(message)
                return
            case "Заново":
                await cmd_start(message)
                return
            case "Доп. информация":
                await cmd_info(message)
                return
            case "Оркестр ассистентов":
                await cmd_band(message)
                return
            case "Использовать оркестр":
                band[message.from_user.id] = True
                await message.answer("Оркестр ассисентов включен! Пришли аудио/видео файл, чтобы начать создание главы", reply_markup=keyboard_start)
                return
            case "Без оркестра":
                band[message.from_user.id] = False
                await message.answer("Оркестр ассисентов отключен! Пришли аудио/видео файл, чтобы начать создание главы", reply_markup=keyboard_start)
                return
            case "Сформировать документ":
                await cmd_doc(message)
                return
            case "Введение":
                await cmd_intro(message)
                return
            case "Заключение":
                await cmd_conclusion(message)
                return
            case "Методика":
                await cmd_metod(message)
                return

        if state[message.from_user.id] == -1:
            dop[message.from_user.id] = message.text
            await message.answer("Спасибо, дополнительная информация учтена! Отправьте аудиофайл для анализа")
            await bot.forward_message(chat_id= reply_chat_id, from_chat_id=message.chat.id, message_id=message.message_id)
            state[message.from_user.id] = 0
            return
        elif state[message.from_user.id] == 0:
            current[message.from_user.id] = message.text
            await message.answer("Текст для создания литературного кейса принят. Нажмите Продолжить для перехода к ассистенту анализа содержания", reply_markup=keyboard_continue)
            await bot.forward_message(chat_id=reply_chat_id, from_chat_id=message.chat.id, message_id=message.message_id)
            state[message.from_user.id] = 1
            return
        gpts[state[message.from_user.id]].add_message(tr, message.text)

        await message.answer(gpts[state[message.from_user.id]].get_answer(tr))
        await bot.forward_message(chat_id= reply_chat_id, from_chat_id=message.chat.id, message_id=message.message_id)
        await bot.forward_message(chat_id= reply_chat_id, from_chat_id=message.chat.id, message_id=message.message_id+1)

    elif message.voice and message.chat.id != reply_chat_id:
        try:
            file=await bot.get_file(message.voice.file_id)
            await process_audio(f"voice_{message.from_user.id}.ogg", file)
            await cmd_continue(message)
        except TelegramBadRequest:
            await message.answer("Файл слишком большой. Разбейте его на несколько глав")
        finally:
            await bot.forward_message(chat_id= reply_chat_id, from_chat_id=message.chat.id, message_id=message.message_id)
            await bot.forward_message(chat_id= reply_chat_id, from_chat_id=message.chat.id, message_id=message.message_id+1)

    elif message.audio and message.chat.id != reply_chat_id:
        try:
            file=await bot.get_file(message.audio.file_id)
            strp = file.file_path.split(".")
            if strp[1] in ('mp3', 'wav', 'ogg', 'oga', 'mkv', 'm4a'):
                await process_audio(f"doc_{message.from_user.id}.{strp[1]}", file)
                if message.caption:
                    current[message.from_user.id] = f"Дополнительная информация: {message.caption}\n" + current[message.from_user.id]
                await cmd_continue(message)
            else:
                await message.answer("Извините, но данный формат аудио не поддерживается.")
        except TelegramBadRequest:
            await message.answer("Файл слишком большой. Разбейте его на несколько глав")
        finally:
            await bot.forward_message(chat_id= reply_chat_id, from_chat_id=message.chat.id, message_id=message.message_id)
            await bot.forward_message(chat_id= reply_chat_id, from_chat_id=message.chat.id, message_id=message.message_id+1)
    elif message.video and message.chat.id != reply_chat_id:
        file=await bot.get_file(message.video.file_id)
        strp = file.file_path.split(".")
        await process_video(file, strp, message.caption)
    elif message.document and message.chat.id != reply_chat_id:
        file=await bot.get_file(message.document.file_id)
        strp = file.file_path.split(".")
        if state[message.from_user.id] not in [-1, -2, -3] and strp[1] in ('mp4', 'avi', 'mkv', 'mov', 'wmv'):
            await process_video(file, strp, message.caption)
            return
        if strp[1] not in ("doc", "docx", "txt", "pdf"):
            await message.answer("Пришлите текстовый файл.")
            return
        text = read_docx_to_string(file.file_path)
        gpt = gpts[state[message.from_user.id]]
        tr = gpt.create_thread()
        gpt.add_message(tr, text)
        answer = gpt.get_answer(tr)
        await send_big_message(bot, message.from_user.id, answer, reply_markup=keyboard_continue)
        await bot.forward_message(chat_id= reply_chat_id, from_chat_id=message.chat.id, message_id=message.message_id)
        await bot.forward_message(chat_id= reply_chat_id, from_chat_id=message.chat.id, message_id=message.message_id+1)
asyncio.run(dp.start_polling(bot))
